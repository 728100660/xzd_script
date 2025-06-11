import markdown
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from spire.doc import *
from spire.doc.common import *


def md_to_pdf_simple(md_text: str, output_path: str = "output.pdf"):
    html = markdown.markdown(md_text)
    c = canvas.Canvas(output_path, pagesize=A4)
    text_obj = c.beginText(40, 800)
    for line in html.splitlines():
        text_obj.textLine(line)
    c.drawText(text_obj)
    c.save()

def md_to_docx_pure(md_text: str, output_file: str = "output.docx"):
    # 转成 HTML
    html = markdown.markdown(md_text)

    # 用 docx 写入
    doc = Document()
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")

    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'li']):
        if element.name.startswith('h'):
            doc.add_heading(element.get_text(), level=int(element.name[1]))
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name == 'li':
            doc.add_paragraph(f"• {element.get_text()}")
    doc.save(output_file)
    return output_file

def md_to_pdf_spire(md_name: str, output_file: str = "output.pdf"):
    # 有水印、要联网、数据外漏
    # 创建一个 Document 类的对象
    doc = Document()

    # 加载一个 Markdown 文件
    doc.LoadFromFile(md_name, FileFormat.Markdown)

    # 将文件保存为 PDF 文档
    doc.SaveToFile(output_file, FileFormat.Docx)

    doc.Dispose()


import pypandoc

def md_to_docx(md_text: str, output_path: str = "output.docx"):
    pypandoc.convert_text(md_text, 'docx', format='md',
                          outputfile=output_path)
    return output_path

def md_to_pdf(md_text: str, output_path: str = "output.pdf"):
    pypandoc.convert_text(md_text, 'pdf', format='md',
                          outputfile=output_path)
    return output_path


def md_to_pdf_m2p(input_file, output_file):
    from markdown2pdf import convert_md_2_pdf
    convert_md_2_pdf(input_file, output_file)


def md_to_pdf_wkhtmltopdf(input, output):
    import pdfkit

    with open(input, encoding='utf-8') as f:
        text = f.read()

    html = markdown.markdown(text, output_format='html', extensions=['tables'])  # MarkDown转HTML


    htmltopdf = r'D:\software\wkhtmltopdf\bin\wkhtmltopdf.exe'
    configuration = pdfkit.configuration(wkhtmltopdf=htmltopdf)
    pdfkit.from_string(html, output_path=output, configuration=configuration, options={'encoding': 'utf-8'})  # HTML转PD


def md_to_pdf_wkhtmltopdf2(md_path, pdf_path):
    from weasyprint import HTML
    # 读取 Markdown 内容
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 转换为 HTML（可以加入 CSS）
    html_content = markdown.markdown(md_content, output_format='html5')

    # 包裹为完整 HTML 文档，含简单 CSS
    html_full = f"""
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            body {{ font-family: "Arial", sans-serif; line-height: 1.6; padding: 40px; }}
            h1, h2, h3 {{ color: #333; }}
            code {{ background-color: #f4f4f4; padding: 2px 4px; border-radius: 4px; }}
            pre {{ background: #f4f4f4; padding: 10px; border-radius: 6px; overflow-x: auto; }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """

    # 转换为 PDF
    HTML(string=html_full).write_pdf(pdf_path)


if __name__ == '__main__':
    md_name = "markdown_test"
    with open(f"{md_name}.md", "r", encoding="utf-8") as f:
        md_text = f.read()

    ## 本地、简单、效果不行
    # md_to_docx_pure(md_text, md_name + ".docx")
    # md_to_pdf_simple(md_text, md_name + ".pdf")

    ## 有水印、要联网、数据外漏
    # md_to_pdf_spire(md_name+".md", md_name + "_spire.docx")


    ## 外部下载环境
    # md_to_docx(md_text, md_name + "_pypandoc.docx")
    # md_to_pdf(md_text, md_name + "_pypandoc.pdf")    # 用不了
    # md_to_pdf_m2p(f"{md_name}.md", md_name + "_m2p.pdf")    # 用不了

    md_to_pdf_wkhtmltopdf(f"{md_name}.md", md_name + "_wkhtmltopdf.pdf")
    # md_to_pdf_wkhtmltopdf2(f"{md_name}.md", md_name + "_wkhtmltopdf2.pdf")